import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont, ImageOps
import base64
import html
import io
import os
from datetime import datetime


PHOTOS_PER_PAGE = 3
A4_WIDTH_PX = 1240
A4_HEIGHT_PX = 1754


def fix_orientation(img):
    try:
        return ImageOps.exif_transpose(img)
    except Exception:
        return img


def img_to_base64(img):
    buf = io.BytesIO()
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    img.save(buf, format="JPEG", quality=80)
    return base64.b64encode(buf.getvalue()).decode()


def _safe_filename(value):
    safe = str(value or "").strip().replace(" ", "_").replace("/", "-").replace("\\", "-")
    return safe or "사진대지"


def _font(size, bold=False):
    candidates = [
        "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf" if bold else "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc" if bold else "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/malgunbd.ttf" if bold else "C:/Windows/Fonts/malgun.ttf",
        "/Library/Fonts/AppleGothic.ttf",
    ]
    for path in candidates:
        if path and os.path.exists(path):
            try:
                return ImageFont.truetype(path, size=size)
            except Exception:
                continue
    return ImageFont.load_default()


def _draw_center(draw, y, text, font, fill):
    bbox = draw.textbbox((0, 0), text, font=font)
    x = (A4_WIDTH_PX - (bbox[2] - bbox[0])) / 2
    draw.text((x, y), text, font=font, fill=fill)


def build_photo_docx(meeting_name, meeting_date, images):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meeting_name.strip())
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(meeting_date.strftime("%Y년 %m월 %d일"))
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    page_width = Inches(6.0)
    for i, (num, name, img) in enumerate(images):
        try:
            export_img = img.convert("RGB") if img.mode in ("RGBA", "P") else img
            buf = io.BytesIO()
            export_img.save(buf, format="JPEG", quality=85)
            buf.seek(0)
            w, h = export_img.size
            aspect = h / w
            img_width = page_width
            img_height = page_width * aspect
            if img_height > Inches(3.5):
                img_height = Inches(3.5)
                img_width = img_height / aspect
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(buf, width=img_width)
            caption = doc.add_paragraph(f"사진 {num}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(9)
            caption.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            doc.add_paragraph()
            if (i + 1) % PHOTOS_PER_PAGE == 0 and (i + 1) < len(images):
                doc.add_page_break()
        except Exception:
            doc.add_paragraph(f"[이미지 로드 실패: {name}]")

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


def _fit_image(img, max_width, max_height):
    export_img = img.convert("RGB") if img.mode in ("RGBA", "P") else img.copy().convert("RGB")
    export_img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
    return export_img


def build_photo_pdf(meeting_name, meeting_date, images):
    title_font = _font(44, bold=True)
    date_font = _font(22)
    caption_font = _font(20)
    pages = []
    grouped = [images[i:i + PHOTOS_PER_PAGE] for i in range(0, len(images), PHOTOS_PER_PAGE)]

    for page_idx, page_images in enumerate(grouped):
        page = Image.new("RGB", (A4_WIDTH_PX, A4_HEIGHT_PX), "white")
        draw = ImageDraw.Draw(page)
        y = 90
        if page_idx == 0:
            _draw_center(draw, y, meeting_name.strip(), title_font, (26, 26, 46))
            y += 62
            _draw_center(draw, y, meeting_date.strftime("%Y년 %m월 %d일"), date_font, (110, 110, 110))
            y += 62

        slot_height = 470 if page_idx == 0 else 520
        for num, name, img in page_images:
            fitted = _fit_image(img, A4_WIDTH_PX - 220, slot_height - 65)
            x = (A4_WIDTH_PX - fitted.width) // 2
            page.paste(fitted, (x, y))
            y += fitted.height + 14
            _draw_center(draw, y, f"사진 {num}", caption_font, (120, 120, 120))
            y += 45
        pages.append(page)

    buffer = io.BytesIO()
    first, rest = pages[0], pages[1:]
    first.save(buffer, format="PDF", save_all=True, append_images=rest, resolution=150.0)
    buffer.seek(0)
    return buffer


def show():
    if st.button("← 홈으로"):
        st.session_state.page = "home"
        st.rerun()

    st.title("📋 회의 사진 대지 생성기")
    st.caption("made by 찰칵혁신단")
    st.markdown("회의 정보를 입력하고 사진을 업로드하면 사진 대지를 자동으로 만들어드려요.")
    st.divider()

    meeting_name = st.text_input("📝 회의명", placeholder="예: 기획팀 월례회의")
    meeting_date = st.date_input("📅 회의 날짜", value=datetime.today())

    uploaded_files = st.file_uploader(
        "📷 사진 업로드 (여러 장 선택 가능)",
        type=["jpg", "jpeg", "png", "heic"],
        accept_multiple_files=True,
    )

    raw_images = []
    if uploaded_files:
        for file in uploaded_files:
            try:
                file.seek(0)
                img = Image.open(io.BytesIO(file.read()))
                img = fix_orientation(img)
                raw_images.append((file.name, img))
            except Exception:
                st.warning(f"'{file.name}' 파일을 읽지 못했습니다.")

    images = []
    if raw_images:
        st.divider()

        if "show_order" not in st.session_state:
            st.session_state.show_order = False
        if st.button("🔢 사진 순서 변경" + (" ▲ 닫기" if st.session_state.show_order else " ▼ 펼치기")):
            st.session_state.show_order = not st.session_state.show_order

        if "order_values" not in st.session_state:
            st.session_state.order_values = {}

        if st.session_state.show_order:
            st.caption("순서를 바꾸고 싶으면 번호를 수정하세요. 같은 번호가 있으면 업로드 순서를 유지합니다.")
            for i, (name, img) in enumerate(raw_images):
                col1, col2 = st.columns([1, 4])
                with col1:
                    order = st.number_input(
                        "순서",
                        min_value=1,
                        max_value=len(raw_images),
                        value=st.session_state.order_values.get(i, i + 1),
                        key=f"order_{i}",
                        label_visibility="collapsed",
                    )
                    st.session_state.order_values[i] = order
                with col2:
                    thumb = img.copy()
                    thumb.thumbnail((400, 200))
                    st.image(thumb, caption=name, width=300)

        order_inputs = []
        for i, (name, img) in enumerate(raw_images):
            order = st.session_state.order_values.get(i, i + 1)
            order_inputs.append((order, i, name, img))
        order_inputs.sort(key=lambda x: (x[0], x[1]))
        images = [(idx + 1, name, img) for idx, (_, _, name, img) in enumerate(order_inputs)]

    if images and meeting_name.strip():
        st.divider()
        st.subheader("📄 다운로드 미리보기")
        date_str = meeting_date.strftime("%Y년 %m월 %d일")
        pages = [images[i:i + PHOTOS_PER_PAGE] for i in range(0, len(images), PHOTOS_PER_PAGE)]

        html_pages = ""
        for page_idx, page_images in enumerate(pages):
            header = ""
            if page_idx == 0:
                header = f"""
                <div style="text-align:center; margin-bottom:6px;">
                    <span style="font-size:22px; font-weight:700; color:#1a1a2e; font-family:'Malgun Gothic',sans-serif;">{html.escape(meeting_name.strip())}</span>
                </div>
                <div style="text-align:center; margin-bottom:24px;">
                    <span style="font-size:12px; color:#888; font-family:'Malgun Gothic',sans-serif;">{date_str}</span>
                </div>
                """
            photos_html = ""
            for num, name, img in page_images:
                b64 = img_to_base64(img)
                photos_html += f"""
                <div style="text-align:center; margin-bottom:20px;">
                    <img src="data:image/jpeg;base64,{b64}" style="max-width:100%; max-height:280px; object-fit:contain; border-radius:4px; box-shadow:0 1px 4px rgba(0,0,0,0.12);">
                    <div style="margin-top:6px; font-size:11px; color:#aaa; font-family:'Malgun Gothic',sans-serif;">사진 {num}</div>
                </div>
                """
            page_num = f'<div style="text-align:right; font-size:10px; color:#ccc; margin-top:4px;">{page_idx + 1} / {len(pages)}</div>'
            separator = '<div style="text-align:center; color:#bbb; margin:12px 0; font-size:11px;">— 페이지 나눔 —</div>' if page_idx < len(pages) - 1 else ""
            html_pages += f'<div style="background:white; border:1px solid #ddd; border-radius:8px; padding:32px 40px 16px 40px; box-shadow:0 2px 8px rgba(0,0,0,0.08); margin-bottom:4px;">{header}{photos_html}{page_num}</div>{separator}'

        components.html(
            f"<html><body style='margin:0;padding:0;background:#f5f5f5;'>{html_pages}</body></html>",
            height=min(len(pages) * (320 * PHOTOS_PER_PAGE + 200), 1200),
            scrolling=True,
        )
    elif uploaded_files and not meeting_name.strip():
        st.info("💡 회의명을 입력하면 문서 미리보기가 표시됩니다.")
    elif meeting_name.strip() and not uploaded_files:
        st.info("💡 사진을 업로드하면 문서 미리보기가 표시됩니다.")

    st.divider()

    if not meeting_name.strip() or not images:
        st.button("📄 파일 생성", type="primary", width="stretch", disabled=True)
        if not meeting_name.strip():
            st.caption("회의명을 입력해야 파일을 생성할 수 있습니다.")
        if not images:
            st.caption("사진을 최소 1장 이상 업로드해야 파일을 생성할 수 있습니다.")
    else:
        if st.button("📄 DOCX 파일 생성", type="primary", width="stretch"):
            with st.spinner("문서를 생성하는 중..."):
                try:
                    doc_buffer = build_photo_docx(meeting_name, meeting_date, images)
                    pdf_buffer = build_photo_pdf(meeting_name, meeting_date, images)

                    date_filename = meeting_date.strftime("%Y%m%d")
                    safe_name = _safe_filename(meeting_name)
                    base_filename = f"{safe_name}_{date_filename}_사진대지"

                    st.success("✅ 문서가 생성되었어요!")
                    # col_docx = st.columns(1)
                    # with col_docx:
                    st.download_button(
                        label="⬇️ DOCX 다운로드",
                        data=doc_buffer,
                        file_name=f"{base_filename}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        width="stretch",
                    )
                    # with col_pdf:
                    #     st.download_button(
                    #         label="⬇️ PDF 다운로드",
                    #         data=pdf_buffer,
                    #         file_name=f"{base_filename}.pdf",
                    #         mime="application/pdf",
                    #         width="stretch",
                    #     )
                except Exception as e:
                    st.error(f"오류: {str(e)}")
                    st.exception(e)
