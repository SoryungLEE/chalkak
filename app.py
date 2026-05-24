import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
import base64
from datetime import datetime

st.set_page_config(
    page_title="회의 사진 대지 생성기",
    page_icon="📋",
    layout="centered"
)

st.title("📋 회의 사진 대지 생성기")
st.caption("made by 찰칵혁신단")
st.markdown("회의 정보를 입력하고 사진을 업로드하면 Word 문서를 자동으로 만들어드려요.")

st.divider()

meeting_name = st.text_input("📝 회의명", placeholder="예: 기획팀 월례회의")
meeting_date = st.date_input("📅 회의 날짜", value=datetime.today())

uploaded_files = st.file_uploader(
    "📷 사진 업로드 (여러 장 선택 가능)",
    type=["jpg", "jpeg", "png", "heic"],
    accept_multiple_files=True
)


def fix_orientation(img):
    try:
        exif_data = img._getexif()
        if exif_data and 274 in exif_data:
            orientation = exif_data[274]
            rotations = {3: 180, 6: 270, 8: 90}
            if orientation in rotations:
                img = img.rotate(rotations[orientation], expand=True)
    except:
        pass
    return img


def img_to_base64(img):
    buf = io.BytesIO()
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    img.save(buf, format="JPEG", quality=80)
    return base64.b64encode(buf.getvalue()).decode()


# 사진 로드
raw_images = []
if uploaded_files:
    for i, file in enumerate(uploaded_files):
        try:
            file.seek(0)
            img = Image.open(io.BytesIO(file.read()))
            img = fix_orientation(img)
            raw_images.append((file.name, img))
        except:
            pass


# 순서 입력
images = []
if raw_images:
    st.divider()
    st.subheader("🔢 사진 순서 설정")
    st.caption("순서를 바꾸고 싶으면 번호를 수정하세요. 같은 번호면 업로드 순서대로 정렬됩니다.")

    order_inputs = []
    for i, (name, img) in enumerate(raw_images):
        col1, col2 = st.columns([1, 4])
        with col1:
            order = st.number_input(
                f"순서",
                min_value=1,
                max_value=len(raw_images),
                value=i + 1,
                key=f"order_{i}",
                label_visibility="collapsed"
            )
        with col2:
            thumb = img.copy()
            thumb.thumbnail((400, 200))
            st.image(thumb, caption=name, use_container_width=False, width=300)
        order_inputs.append((order, i, name, img))

    # 순서대로 정렬
    order_inputs.sort(key=lambda x: (x[0], x[1]))
    images = [(idx + 1, name, img) for idx, (_, _, name, img) in enumerate(order_inputs)]


# 문서 미리보기
if images and meeting_name.strip():
    st.divider()
    st.subheader("📄 문서 미리보기")

    date_str = meeting_date.strftime("%Y년 %m월 %d일")
    PHOTOS_PER_PAGE = 3
    pages = [images[i:i+PHOTOS_PER_PAGE] for i in range(0, len(images), PHOTOS_PER_PAGE)]

    html_pages = ""
    for page_idx, page_images in enumerate(pages):
        header = ""
        if page_idx == 0:
            header = f"""
            <div style="text-align:center; margin-bottom:6px;">
                <span style="font-size:22px; font-weight:700; color:#1a1a2e; font-family:'Malgun Gothic',sans-serif;">
                    {meeting_name.strip()}
                </span>
            </div>
            <div style="text-align:center; margin-bottom:24px;">
                <span style="font-size:12px; color:#888; font-family:'Malgun Gothic',sans-serif;">
                    {date_str}
                </span>
            </div>
            """

        photos_html = ""
        for num, name, img in page_images:
            b64 = img_to_base64(img)
            photos_html += f"""
            <div style="text-align:center; margin-bottom:20px;">
                <img src="data:image/jpeg;base64,{b64}"
                     style="max-width:100%; max-height:280px; object-fit:contain; border-radius:4px; box-shadow:0 1px 4px rgba(0,0,0,0.12);">
                <div style="margin-top:6px; font-size:11px; color:#aaa; font-family:'Malgun Gothic',sans-serif;">사진 {num}</div>
            </div>
            """

        page_num = f"""
        <div style="text-align:right; font-size:10px; color:#ccc; margin-top:4px; font-family:'Malgun Gothic',sans-serif;">
            {page_idx+1} / {len(pages)}
        </div>
        """

        separator = ""
        if page_idx < len(pages) - 1:
            separator = """<div style="text-align:center; color:#bbb; margin:12px 0; font-size:11px;">— 페이지 나눔 —</div>"""

        html_pages += f"""
        <div style="background:white; border:1px solid #ddd; border-radius:8px; padding:32px 40px 16px 40px; box-shadow:0 2px 8px rgba(0,0,0,0.08); margin-bottom:4px;">
            {header}{photos_html}{page_num}
        </div>
        {separator}
        """

    full_html = f"<html><body style='margin:0;padding:0;background:#f5f5f5;'>{html_pages}</body></html>"
    total_height = len(pages) * (320 * PHOTOS_PER_PAGE + 200)
    components.html(full_html, height=min(total_height, 1200), scrolling=True)

elif uploaded_files and not meeting_name.strip():
    st.info("💡 회의명을 입력하면 문서 미리보기가 표시됩니다.")
elif meeting_name.strip() and not uploaded_files:
    st.info("💡 사진을 업로드하면 문서 미리보기가 표시됩니다.")


# 문서 생성
st.divider()

if st.button("📄 Word 문서 생성", type="primary", use_container_width=True):
    if not meeting_name.strip():
        st.error("⚠️ 회의명을 입력해주세요!")
    elif not images:
        st.error("⚠️ 사진을 최소 1장 이상 업로드해주세요!")
    else:
        with st.spinner("문서를 생성하는 중..."):
            try:
                doc = Document()
                section = doc.sections[0]
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

                # 첫 페이지 헤더
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = title.add_run(meeting_name.strip())
                run.bold = True
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

                date_p = doc.add_paragraph()
                date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                date_run = date_p.add_run(meeting_date.strftime("%Y년 %m월 %d일"))
                date_run.font.size = Pt(11)
                date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                doc.add_paragraph()

                PHOTOS_PER_PAGE = 3
                page_width = Inches(6.0)

                for i, (num, name, img) in enumerate(images):
                    try:
                        if img.mode in ("RGBA", "P"):
                            img = img.convert("RGB")

                        buf = io.BytesIO()
                        img.save(buf, format="JPEG", quality=85)
                        buf.seek(0)

                        w, h = img.size
                        aspect = h / w
                        img_width = page_width
                        img_height = page_width * aspect
                        if img_height > Inches(3.5):
                            img_height = Inches(3.5)
                            img_width = img_height / aspect

                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(buf, width=img_width)

                        caption = doc.add_paragraph(f"사진 {num}")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.runs[0].font.size = Pt(9)
                        caption.runs[0].font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
                        doc.add_paragraph()

                        if (i + 1) % PHOTOS_PER_PAGE == 0 and (i + 1) < len(images):
                            doc.add_page_break()

                    except:
                        doc.add_paragraph(f"[이미지 로드 실패: {name}]")

                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                date_filename = meeting_date.strftime("%Y%m%d")
                safe_name = meeting_name.strip().replace(" ", "_").replace("/", "-")
                filename = f"{safe_name}_{date_filename}_사진대지.docx"

                st.success("✅ 문서가 생성되었어요!")
                st.download_button(
                    label="⬇️ Word 문서 다운로드",
                    data=doc_buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            except Exception as e:
                st.error(f"오류: {str(e)}")
                st.exception(e)
