import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageOps, ImageEnhance, ImageFilter
import io
import base64
import json
import requests
import re
import os
import subprocess
import tempfile
import zipfile
import uuid
from datetime import datetime
from xml.sax.saxutils import escape


def fix_orientation(img):
    """EXIF 회전값을 실제 픽셀에 반영합니다."""
    try:
        return ImageOps.exif_transpose(img)
    except Exception:
        return img


def _resize_for_vision(img, min_width=1200, max_side=3500):
    """
    영수증 사진은 폭이 너무 작으면 품명/수량 열을 모델이 자주 틀립니다.
    311px 같은 작은 이미지는 먼저 키우고, 너무 큰 이미지는 API 비용/오류를 막기 위해 줄입니다.
    """
    w, h = img.size

    if w < min_width:
        ratio = min_width / w
        img = img.resize((int(w * ratio), int(h * ratio)), Image.Resampling.LANCZOS)

    w, h = img.size
    if max(w, h) > max_side:
        ratio = max_side / max(w, h)
        img = img.resize((int(w * ratio), int(h * ratio)), Image.Resampling.LANCZOS)

    return img


def make_receipt_vision_images(img):
    """
    모델에 2장을 보냅니다.
    1) 원본 보정본: 색/배경 맥락 유지
    2) 흑백 강조본: 작은 글씨/숫자열 판독 강화
    """
    img = fix_orientation(img).convert("RGB")
    color = _resize_for_vision(img)

    gray = ImageOps.grayscale(color)
    gray = ImageOps.autocontrast(gray, cutoff=1)
    gray = ImageEnhance.Contrast(gray).enhance(1.8)
    gray = gray.filter(ImageFilter.UnsharpMask(radius=1, percent=170, threshold=3))
    gray = gray.convert("RGB")

    return color, gray


def image_to_data_url(img, fmt="JPEG", quality=95):
    buf = io.BytesIO()
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    img.save(buf, format=fmt, quality=quality, optimize=True)
    mime = "image/jpeg" if fmt.upper() == "JPEG" else "image/png"
    return f"data:{mime};base64,{base64.b64encode(buf.getvalue()).decode('utf-8')}"


# 기존 다른 코드가 img_to_base64를 호출해도 깨지지 않게 유지
# 단, read_receipt_with_ai에서는 data URL을 직접 사용합니다.
def img_to_base64(img):
    color, _ = make_receipt_vision_images(img)
    buf = io.BytesIO()
    color.save(buf, format="JPEG", quality=95, optimize=True)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


RECEIPT_SCHEMA = {
    "name": "korean_receipt_ocr",
    "strict": True,
    "schema": {
        "type": "object",
        "properties": {
            "store": {"type": "string"},
            "date": {"type": "string"},
            "lines": {
                "type": "array",
                "items": {"type": "string"}
            },
            "items": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "raw_line": {"type": "string"},
                        "name": {"type": "string"},
                        "spec": {"type": "string"},
                        "quantity": {"type": "number"},
                        "unit_price": {"type": "number"},
                        "amount": {"type": "number"}
                    },
                    "required": ["raw_line", "name", "spec", "quantity", "unit_price", "amount"],
                    "additionalProperties": False
                }
            },
            "warnings": {
                "type": "array",
                "items": {"type": "string"}
            }
        },
        "required": ["store", "date", "lines", "items", "warnings"],
        "additionalProperties": False
    }
}


RECEIPT_PROMPT = """너는 한국 영수증 OCR 엔진이다. 최종 판단보다 원문 보존이 더 중요하다.

해야 할 일:
1. 영수증에서 보이는 텍스트 줄을 위에서 아래 순서대로 lines에 최대한 그대로 넣어라.
2. 상품 구매로 보이는 줄은 items에도 넣어라.
3. 상품행은 raw_line에 한 줄 원문을 반드시 넣어라.
4. 상품명 안의 숫자, 괄호, 용량, 규격, 세트표기는 절대 떼지 마라.
   예: 25P, 21*21cm, 500ml, 2입, A4, 80g, 3개입은 상품명 일부다.
5. 단가/수량/금액이 보이면 그대로 넣되, 확실하지 않으면 raw_line 중심으로 넣어라.
6. 날짜는 YYYY-MM-DD로 가능하면 넣고, 아니면 빈 문자열.
7. JSON 외 텍스트를 절대 쓰지 마라.

주의:
- 이마트/다이소/편의점/식당/온라인몰/카드전표 등 형식이 다를 수 있다.
- 오른쪽 숫자열이 [단가 수량 금액]일 수도 있고, [수량 단가 금액]일 수도 있고, 금액만 있을 수도 있다.
- Python 후처리가 최종 검증을 하므로 lines와 raw_line을 최대한 정확히 보존하는 것이 최우선이다.
"""


_AMOUNT_TOKEN = r"(?:\d{1,3}(?:,\d{3})+|\d+)"
_NUM_RE = re.compile(_AMOUNT_TOKEN)

# 진짜 합계/결제/세금 줄만 제외. '봉투'는 실제 구매품일 수 있어서 제외하지 않는다.
_EXCLUDE_NAME_RE = re.compile(
    r"(과세\s*합계|부가\s*세|판매\s*합계|총\s*합계|합\s*계|받은\s*금액|거스름|"
    r"신용\s*카드|체크\s*카드|현금\s*영수증|승인|카드번호|포인트|멤버십|"
    r"사업자|대표|전화|주소|POS|교환|환불|영수증|바코드)",
    re.IGNORECASE,
)
_CODE_ONLY_RE = re.compile(r"^\s*\[?\s*\d{4,}\s*\]?\s*$")


def _clean_space(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _to_int(value):
    if isinstance(value, bool):
        return 0
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return int(round(value))
    s = re.sub(r"[^0-9]", "", str(value or ""))
    return int(s) if s else 0


def _normalize_date(value):
    s = str(value or "")
    m = re.search(r"(20\d{2})\D+(\d{1,2})\D+(\d{1,2})", s)
    if m:
        y, mo, d = map(int, m.groups())
        return f"{y:04d}-{mo:02d}-{d:02d}"
    m = re.search(r"(\d{2})[./-](\d{1,2})[./-](\d{1,2})", s)
    if m:
        yy, mo, d = map(int, m.groups())
        return f"20{yy:02d}-{mo:02d}-{d:02d}"
    return ""


def _strip_code_lines(text):
    # 상품코드 [1002504] 같은 토큰만 제거. 상품명 내부 숫자는 보존.
    text = re.sub(r"\s*\[\s*\d{4,}\s*\]\s*", " ", str(text or ""))
    return _clean_space(text.replace("원", ""))


def _looks_like_non_item_name(name):
    name = _clean_space(name)
    if not name:
        return True
    if _CODE_ONLY_RE.match(name):
        return True
    if _EXCLUDE_NAME_RE.search(name):
        return True
    if not re.search(r"[가-힣A-Za-z]", name):
        return True
    return False


def _parse_item_raw_line(raw_line):
    """
    영수증 종류별 숫자 패턴을 Python이 최종 판정한다.
    지원:
    - 다이소형: 품명 단가 수량 금액
    - 마트형: 품명 수량 단가 금액
    - 일반형: 품명 금액  -> 수량 1, 단가=금액
    - x 표기형: 품명 2 x 1,000 2,000
    """
    raw = _strip_code_lines(raw_line)
    if not raw or _EXCLUDE_NAME_RE.search(raw):
        return None

    # 할인/취소/마이너스 줄은 일단 제외
    if re.search(r"(할인|에누리|쿠폰|취소|반품)", raw):
        return None

    # 1) x 표기: 품명 2 x 1,000 2,000 또는 품명 1,000 x 2 2,000
    m = re.match(rf"^(.+?)\s+({_AMOUNT_TOKEN})\s*[xX×]\s*({_AMOUNT_TOKEN})\s+({_AMOUNT_TOKEN})\s*$", raw)
    if m:
        name = _clean_space(m.group(1))
        a, b, total = _to_int(m.group(2)), _to_int(m.group(3)), _to_int(m.group(4))
        candidates = []
        if a * b == total:
            # 보통 작은 숫자는 수량, 큰 숫자는 단가
            if a <= 99 and b > 99:
                candidates.append((b, a, total))
            elif b <= 99 and a > 99:
                candidates.append((a, b, total))
            else:
                candidates.append((b, a, total))
        if candidates and not _looks_like_non_item_name(name):
            unit_price, quantity, amount = candidates[0]
            return {"raw_line": raw, "name": name, "spec": "", "quantity": quantity, "unit_price": unit_price, "amount": amount}

    nums = list(_NUM_RE.finditer(raw))
    if not nums:
        return None

    # 2) 오른쪽 끝 숫자 3개에서 산술검증. 품명 내부 숫자는 앞쪽에 남긴다.
    if len(nums) >= 3:
        best = None
        for i in range(len(nums) - 3, -1, -1):
            n1 = _to_int(nums[i].group())
            n2 = _to_int(nums[i + 1].group())
            n3 = _to_int(nums[i + 2].group())
            if n1 <= 0 or n2 <= 0 or n3 <= 0:
                continue

            name = _clean_space(raw[:nums[i].start()])
            if _looks_like_non_item_name(name):
                continue

            # 단가 수량 금액
            if n1 * n2 == n3:
                best = (name, n2, n1, n3)
                break
            # 수량 단가 금액
            if n2 * n1 == n3:
                best = (name, n1, n2, n3)
                break

        if best:
            name, quantity, unit_price, amount = best
            return {"raw_line": raw, "name": name, "spec": "", "quantity": quantity, "unit_price": unit_price, "amount": amount}

    # 3) 금액만 있는 상품행: 품명 금액 -> 수량 1
    last = nums[-1]
    amount = _to_int(last.group())
    name = _clean_space(raw[:last.start()])
    if amount > 0 and not _looks_like_non_item_name(name):
        # 너무 작은 숫자만 있으면 코드/수량일 가능성이 높음
        if amount >= 100:
            return {"raw_line": raw, "name": name, "spec": "", "quantity": 1, "unit_price": amount, "amount": amount}

    return None


def _merge_wrapped_lines(lines):
    """온라인몰/마트처럼 상품명이 다음 줄로 이어지는 경우를 조금 더 버티게 합친다."""
    cleaned = [_strip_code_lines(x) for x in lines or []]
    cleaned = [x for x in cleaned if x and not _CODE_ONLY_RE.match(x)]

    merged = []
    pending = ""
    for line in cleaned:
        if _EXCLUDE_NAME_RE.search(line):
            if pending:
                merged.append(pending)
                pending = ""
            continue

        parsed = _parse_item_raw_line(line)
        if parsed:
            if pending and not _parse_item_raw_line(pending):
                combined = _clean_space(pending + " " + line)
                if _parse_item_raw_line(combined):
                    merged.append(combined)
                    pending = ""
                    continue
                merged.append(pending)
                pending = ""
            merged.append(line)
            continue

        # 숫자가 거의 없고 글자가 있으면 다음 줄 상품명 후보로 보관
        if re.search(r"[가-힣A-Za-z]", line):
            if pending:
                merged.append(pending)
            pending = line

    if pending:
        merged.append(pending)
    return merged


def _normalize_receipt_result(data):
    warnings = list(data.get("warnings") or [])
    normalized = {
        "store": _clean_space(data.get("store", "")),
        "date": _normalize_date(data.get("date", "")),
        "items": [],
        "warnings": warnings,
    }

    candidates = []
    for line in data.get("lines", []) or []:
        candidates.append(line)
    for item in data.get("items", []) or []:
        candidates.append(item.get("raw_line") or "")

    for line in _merge_wrapped_lines(candidates):
        parsed = _parse_item_raw_line(line)
        if parsed:
            normalized["items"].append(parsed)

    # raw_line에서 못 잡힌 것만 AI 필드를 보조로 사용
    for item in data.get("items", []) or []:
        raw_line = item.get("raw_line", "")
        name = _clean_space(item.get("name", ""))
        unit_price = _to_int(item.get("unit_price"))
        quantity = _to_int(item.get("quantity"))
        amount = _to_int(item.get("amount"))

        if _looks_like_non_item_name(name):
            continue
        if amount <= 0:
            continue
        if quantity <= 0:
            quantity = 1
        if unit_price <= 0:
            unit_price = amount // quantity if quantity else amount
        if unit_price * quantity != amount:
            # OCR이 금액만 정확한 경우는 1개 구매로 살림
            if quantity == 1:
                unit_price = amount
            else:
                normalized["warnings"].append(f"금액 검증 실패로 제외: {raw_line or name}")
                continue

        normalized["items"].append({
            "raw_line": _clean_space(raw_line),
            "name": name,
            "spec": "",
            "quantity": quantity,
            "unit_price": unit_price,
            "amount": amount,
        })

    # 중복 제거
    deduped = []
    seen = set()
    for it in normalized["items"]:
        key = (it["name"], it["quantity"], it["unit_price"], it["amount"])
        if key in seen:
            continue
        seen.add(key)
        deduped.append(it)
    normalized["items"] = deduped

    return normalized


def read_receipt_with_ai(img):
    color_img, contrast_img = make_receipt_vision_images(img)
    color_url = image_to_data_url(color_img)
    contrast_url = image_to_data_url(contrast_img)
    api_key = _get_openai_key()
    if not api_key:
        return None, "OPENAI_API_KEY가 설정되어 있지 않습니다."

    payload = {
        "model": "gpt-4o-2024-08-06",
        "temperature": 0,
        "max_tokens": 3000,
        "response_format": {
            "type": "json_schema",
            "json_schema": RECEIPT_SCHEMA,
        },
        "messages": [
            {
                "role": "system",
                "content": "너는 한국 영수증 OCR 엔진이다. 원문 줄 보존을 최우선으로 하고, JSON만 반환한다.",
            },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": RECEIPT_PROMPT + "\n첫 번째 이미지는 원본 보정본, 두 번째 이미지는 흑백 대비 강화본이다. 둘을 대조해서 lines를 최대한 정확히 작성해라."},
                    {"type": "image_url", "image_url": {"url": color_url, "detail": "high"}},
                    {"type": "image_url", "image_url": {"url": contrast_url, "detail": "high"}},
                ],
            },
        ],
    }

    try:
        res = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            },
            json=payload,
            timeout=60,
        )
        res.raise_for_status()
        res_json = res.json()

        choice = (res_json.get("choices") or [{}])[0]
        message = choice.get("message", {})
        if message.get("refusal"):
            return None, f"모델 거절: {message['refusal']}"

        text = (message.get("content") or "").strip()
        if not text:
            return None, f"빈 응답: {str(res_json)[:300]}"

        parsed = json.loads(text)
        parsed = _normalize_receipt_result(parsed)
        if not parsed["items"]:
            return None, "상품행을 검증하지 못했습니다. 사진을 더 가까이/선명하게 찍거나 수동 입력해 주세요."
        return parsed, None

    except requests.exceptions.HTTPError as e:
        try:
            detail = res.json()
        except Exception:
            detail = res.text
        return None, f"API HTTP 오류: {e} / {str(detail)[:500]}"
    except Exception as e:
        return None, str(e)


DEFAULT_INSPECTION_ITEM = "물품상태 및 수량 등"
DEFAULT_INSPECTION_RESULT = "이상없음"


def _safe_filename(value):
    safe = str(value or "").strip().replace(" ", "_").replace("/", "-").replace("\\", "-")
    return safe or "물품보고서"


def _new_receipt_row_id():
    return uuid.uuid4().hex


def _ensure_receipt_row_ids(items):
    for item in items:
        if not item.get("_row_id"):
            item["_row_id"] = _new_receipt_row_id()
    return items


def _get_openai_key():
    try:
        return st.secrets.get("OPENAI_API_KEY", "")
    except Exception:
        return ""


def read_receipt_with_tesseract(img):
    """OpenAI 크레딧이 없거나 API가 실패할 때 사용할 로컬 OCR 보조 수단입니다."""
    color_img, contrast_img = make_receipt_vision_images(img)
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp_path = tmp.name
            contrast_img.save(tmp, format="PNG")
        cmd = ["tesseract", tmp_path, "stdout", "-l", "kor+eng", "--psm", "6"]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60, check=False)
        if proc.returncode != 0:
            cmd = ["tesseract", tmp_path, "stdout", "-l", "eng", "--psm", "6"]
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60, check=False)
        if proc.returncode != 0:
            return None, f"로컬 Tesseract OCR 실행 실패: {proc.stderr.strip()}"
        raw_text = proc.stdout
    except FileNotFoundError:
        return None, "로컬 Tesseract OCR을 사용할 수 없습니다. 서버에 tesseract와 kor 언어팩을 설치해야 합니다."
    except Exception as e:
        return None, f"로컬 Tesseract OCR 오류: {e}"
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    lines = [_clean_space(line) for line in raw_text.splitlines() if _clean_space(line)]
    parsed = _normalize_receipt_result({"store": "", "date": "", "lines": lines, "items": [], "warnings": []})

    for line in lines[:8]:
        if not parsed["date"]:
            parsed["date"] = _normalize_date(line)
        if not parsed["store"] and re.search(r"[가-힣A-Za-z]", line) and not _EXCLUDE_NAME_RE.search(line):
            parsed["store"] = line

    parsed["warnings"].append("로컬 OCR 결과입니다. OpenAI Vision보다 정확도가 낮을 수 있어 반드시 직접 확인해야 합니다.")
    if not parsed["items"]:
        return None, "로컬 OCR에서도 상품행을 검증하지 못했습니다. 수동으로 항목을 추가해 주세요."
    return parsed, None


def read_receipt(img):
    api_key = _get_openai_key()
    if api_key:
        result, err = read_receipt_with_ai(img)
        if result:
            return result, None, "OpenAI Vision"
        fallback, fallback_err = read_receipt_with_tesseract(img)
        if fallback:
            fallback["warnings"].insert(0, f"OpenAI 분석 실패 후 로컬 OCR로 전환했습니다. OpenAI 오류: {err}")
            return fallback, None, "Local Tesseract OCR"
        return None, f"OpenAI 오류: {err}\n로컬 OCR 오류: {fallback_err}", ""

    result, err = read_receipt_with_tesseract(img)
    return result, err, "Local Tesseract OCR" if result else ""


def _receipt_headers():
    return ["물품명", "규격", "단위", "수량", "금액", "검사항목", "검사결과", "검사일자"]


def build_receipt_docx(report_title, report_date, purchase_date, store_name, items):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run(report_title.strip())
    t_run.bold = True
    t_run.font.size = Pt(18)
    t_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_run = date_p.add_run(report_date.strftime("%Y년 %m월 %d일"))
    d_run.font.size = Pt(11)
    d_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    def add_info_row(label, value):
        p = doc.add_paragraph()
        label_run = p.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.size = Pt(11)
        val_run = p.add_run(str(value or ""))
        val_run.font.size = Pt(11)

    add_info_row("상호명", store_name)
    add_info_row("구매날짜", purchase_date.strftime("%Y년 %m월 %d일"))
    doc.add_paragraph()

    headers = _receipt_headers()
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E8EDF5")
        tcPr.append(shd)

    for item in items:
        row = table.add_row()
        vals = [item.get("물품명", ""), item.get("규격", ""), item.get("단위", ""), str(item.get("수량", "")), f"{int(item.get('금액', 0)):,}", item.get("검사항목", ""), item.get("검사결과", ""), item.get("검사일자", "")]
        for j, (v, a) in enumerate(zip(vals, aligns)):
            cell = row.cells[j]
            cell.paragraphs[0].alignment = a
            run = cell.paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


def _xlsx_col_name(index):
    name = ""
    while index:
        index, rem = divmod(index - 1, 26)
        name = chr(65 + rem) + name
    return name


def _xlsx_cell(row, col, value, style=None):
    ref = f"{_xlsx_col_name(col)}{row}"
    style_attr = f' s="{style}"' if style is not None else ""
    if isinstance(value, (int, float)):
        return f'<c r="{ref}"{style_attr}><v>{value}</v></c>'
    return f'<c r="{ref}" t="inlineStr"{style_attr}><is><t>{escape(str(value or ""))}</t></is></c>'


def build_receipt_xlsx(report_title, report_date, purchase_date, store_name, items):
    headers = _receipt_headers()
    rows = []
    rows.append(f'<row r="1">{_xlsx_cell(1, 1, report_title.strip(), 1)}</row>')
    rows.append(f'<row r="2">{_xlsx_cell(2, 1, "보고서 날짜")}{_xlsx_cell(2, 2, report_date.strftime("%Y-%m-%d"))}</row>')
    rows.append(f'<row r="3">{_xlsx_cell(3, 1, "상호명")}{_xlsx_cell(3, 2, store_name)}</row>')
    rows.append(f'<row r="4">{_xlsx_cell(4, 1, "구매날짜")}{_xlsx_cell(4, 2, purchase_date.strftime("%Y-%m-%d"))}</row>')
    rows.append(f'<row r="6">{"".join(_xlsx_cell(6, i, h, 2) for i, h in enumerate(headers, start=1))}</row>')
    for r, item in enumerate(items, start=7):
        values = [item.get("물품명", ""), item.get("규격", ""), item.get("단위", ""), item.get("수량", ""), item.get("금액", 0), item.get("검사항목", ""), item.get("검사결과", ""), item.get("검사일자", "")]
        rows.append(f'<row r="{r}">{"".join(_xlsx_cell(r, c, v) for c, v in enumerate(values, start=1))}</row>')

    sheet = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<cols><col min="1" max="1" width="28" customWidth="1"/><col min="2" max="8" width="16" customWidth="1"/></cols>
<sheetData>{''.join(rows)}</sheetData><mergeCells count="1"><mergeCell ref="A1:H1"/></mergeCells></worksheet>'''
    styles = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"><fonts count="3"><font/><font><b/><sz val="16"/></font><font><b/></font></fonts><fills count="3"><fill><patternFill patternType="none"/></fill><fill><patternFill patternType="gray125"/></fill><fill><patternFill patternType="solid"><fgColor rgb="FFE8EDF5"/><bgColor indexed="64"/></patternFill></fill></fills><borders count="1"><border/></borders><cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs><cellXfs count="3"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/><xf numFmtId="0" fontId="1" fillId="0" borderId="0" xfId="0" applyFont="1"/><xf numFmtId="0" fontId="2" fillId="2" borderId="0" xfId="0" applyFont="1" applyFill="1"/></cellXfs></styleSheet>'''
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/><Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/><Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/></Types>''')
        zf.writestr("_rels/.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/></Relationships>''')
        zf.writestr("xl/workbook.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><sheets><sheet name="물품보고서" sheetId="1" r:id="rId1"/></sheets></workbook>''')
        zf.writestr("xl/_rels/workbook.xml.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/></Relationships>''')
        zf.writestr("xl/worksheets/sheet1.xml", sheet)
        zf.writestr("xl/styles.xml", styles)
    buffer.seek(0)
    return buffer


def show():
    if st.button("← 홈으로"):
        st.session_state.page = "home"
        st.session_state.receipt_items = []
        st.session_state.last_uploaded = None
        st.session_state.analyzed = False
        st.session_state.confirmed = False
        st.rerun()

    st.title("🧾 영수증 물품보고서 생성기")
    st.caption("made by 찰칵혁신단")
    st.markdown("영수증 사진을 올리면 OCR이 항목을 읽고, 사람이 확인·수정한 뒤 엑셀로 다운로드할 수 있습니다.")
    st.info("ℹ️ OpenAI API 크레딧이 없거나 API가 실패하면 로컬 Tesseract OCR을 시도합니다. 로컬 OCR은 서버에 Tesseract와 한국어 언어팩이 설치되어 있어야 하며 정확도가 낮을 수 있습니다.")
    st.divider()

    report_title = st.text_input("📝 보고서 제목", placeholder="예: 2026년 6월 사무용품 물품보고서")
    report_date = st.date_input("📅 보고서 날짜", value=datetime.today())
    st.divider()

    uploaded_file = st.file_uploader(
        "📷 영수증 사진 업로드 (1장)",
        type=["jpg", "jpeg", "png", "heic"],
        accept_multiple_files=False,
    )

    if "receipt_items" not in st.session_state:
        st.session_state.receipt_items = []
    if "receipt_meta" not in st.session_state:
        st.session_state.receipt_meta = {"store": "", "date": ""}
    if "last_uploaded" not in st.session_state:
        st.session_state.last_uploaded = None
    if "analyzed" not in st.session_state:
        st.session_state.analyzed = False
    if "confirmed" not in st.session_state:
        st.session_state.confirmed = False
    if "receipt_warnings" not in st.session_state:
        st.session_state.receipt_warnings = []
    if "ocr_provider" not in st.session_state:
        st.session_state.ocr_provider = ""
    if "img_rotation" not in st.session_state:
        st.session_state.img_rotation = 0

    current_name = uploaded_file.name if uploaded_file else None
    if current_name != st.session_state.last_uploaded:
        st.session_state.receipt_items = []
        st.session_state.receipt_meta = {"store": "", "date": ""}
        st.session_state.analyzed = False
        st.session_state.confirmed = False
        st.session_state.receipt_warnings = []
        st.session_state.ocr_provider = ""
        st.session_state.img_rotation = 0
        st.session_state.last_uploaded = current_name

    if uploaded_file:
        uploaded_file.seek(0)
        img_orig = Image.open(io.BytesIO(uploaded_file.read()))
        img = fix_orientation(img_orig)
        rotation = st.session_state.img_rotation
        if rotation != 0:
            img = img.rotate(-rotation, expand=True)

        st.warning("⚠️ AI/OCR이 읽은 내용은 틀릴 수 있습니다. 다운로드 전 상호명, 구매날짜, 물품명, 수량, 금액을 반드시 직접 확인하세요.")
        col_img, col_btn = st.columns([2, 1])
        with col_img:
            st.image(img, caption=uploaded_file.name, width="stretch")
        with col_btn:
            st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
            if st.button("🔄 90도 회전", width="stretch"):
                st.session_state.img_rotation = (st.session_state.img_rotation + 90) % 360
                st.rerun()
            st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
            if st.button("🤖 영수증 읽기", type="primary", width="stretch"):
                with st.spinner("OCR이 영수증을 읽는 중..."):
                    result, err, provider = read_receipt(img)
                    if result and result.get("items"):
                        st.session_state.receipt_meta = {
                            "store": result.get("store", ""),
                            "date": result.get("date", ""),
                        }
                        today = datetime.today().strftime("%Y-%m-%d")
                        st.session_state.receipt_items = [
                            {
                                "_row_id": _new_receipt_row_id(),
                                "물품명": item.get("name", ""),
                                "규격": item.get("spec", "") or "-",
                                "단위": "개",
                                "수량": item.get("quantity", 1),
                                "금액": item.get("amount", 0),
                                "검사항목": DEFAULT_INSPECTION_ITEM,
                                "검사결과": DEFAULT_INSPECTION_RESULT,
                                "검사일자": today,
                            }
                            for item in result["items"]
                        ]
                        st.session_state.analyzed = True
                        st.session_state.confirmed = False
                        st.session_state.receipt_warnings = result.get("warnings", [])
                        st.session_state.ocr_provider = provider
                        st.success(f"✅ {provider}로 {len(st.session_state.receipt_items)}개 항목을 읽었어요!")
                        if st.session_state.receipt_warnings:
                            with st.expander("OCR 경고 보기", expanded=True):
                                for warning in st.session_state.receipt_warnings:
                                    st.warning(warning)
                    else:
                        st.error("❌ 항목을 읽지 못했어요. 아래 표에서 수동으로 입력할 수 있도록 먼저 영수증을 더 선명하게 찍어 다시 시도해 주세요.")
                        if err:
                            st.caption(f"오류 상세: {err}")

    if st.session_state.analyzed and st.session_state.receipt_items:
        st.divider()
        st.subheader("📋 기본 정보")
        st.caption(f"읽기 방식: {st.session_state.ocr_provider or '알 수 없음'}")
        st.warning("⚠️ AI가 읽었으므로 반드시 확인이 필요합니다.")

        meta = st.session_state.receipt_meta
        try:
            meta_date = datetime.strptime(meta["date"], "%Y-%m-%d").date()
        except Exception:
            meta_date = datetime.today().date()

        locked = st.session_state.confirmed
        col_a, col_b = st.columns(2)
        with col_a:
            store_name = st.text_input("상호명", value=meta.get("store", ""), key="meta_store", disabled=locked)
        with col_b:
            purchase_date = st.date_input("구매날짜", value=meta_date, key="meta_date", disabled=locked)

        st.session_state.receipt_meta = {
            "store": store_name,
            "date": purchase_date.strftime("%Y-%m-%d"),
        }

        st.divider()
        if locked:
            st.success("✅ 이상없음을 확인했습니다. 수정하려면 '잠금 해제'를 눌러주세요.")
        else:
            st.subheader("📝 항목 확인 및 수정")
            st.caption("없는 내용은 만들지 말고 빈칸 또는 '-'로 두세요. 잘못 읽은 숫자는 영수증과 대조해 직접 수정하세요.")

        cols_w = [2.4, 1.4, 0.8, 0.8, 1.3, 2.3, 1.4, 1.3, 0.6]
        h_labels = ["물품명", "규격", "단위", "수량", "금액", "검사항목", "검사결과", "검사일자", ""]
        hcols = st.columns(cols_w)
        for hc, hl in zip(hcols, h_labels):
            hc.markdown(f"**{hl}**")

        _ensure_receipt_row_ids(st.session_state.receipt_items)
        updated_items = []
        to_delete_ids = set()
        for i, item in enumerate(st.session_state.receipt_items):
            row_id = item["_row_id"]
            c = st.columns(cols_w)
            with c[0]:
                name = st.text_input("물품명", value=item.get("물품명", ""), key=f"name_{row_id}", label_visibility="collapsed", disabled=locked)
            with c[1]:
                spec = st.text_input("규격", value=item.get("규격", "-"), key=f"spec_{row_id}", label_visibility="collapsed", disabled=locked)
            with c[2]:
                unit = st.text_input("단위", value=item.get("단위", "개"), key=f"unit_{row_id}", label_visibility="collapsed", disabled=locked)
            with c[3]:
                qty = st.number_input("수량", value=int(item.get("수량", 1)), min_value=1, key=f"qty_{row_id}", label_visibility="collapsed", disabled=locked)
            with c[4]:
                amount = st.number_input("금액", value=int(item.get("금액", 0)), min_value=0, key=f"amount_{row_id}", label_visibility="collapsed", disabled=locked)
            with c[5]:
                insp_item = st.text_input("검사항목", value=item.get("검사항목", DEFAULT_INSPECTION_ITEM), key=f"insp_{row_id}", label_visibility="collapsed", disabled=locked)
            with c[6]:
                insp_result = st.text_input("검사결과", value=item.get("검사결과", DEFAULT_INSPECTION_RESULT), key=f"iresult_{row_id}", label_visibility="collapsed", disabled=locked)
            with c[7]:
                try:
                    insp_value = datetime.strptime(item.get("검사일자", ""), "%Y-%m-%d").date()
                except Exception:
                    insp_value = datetime.today().date()
                insp_date = st.date_input("검사일자", value=insp_value, key=f"idate_{row_id}", label_visibility="collapsed", disabled=locked)
            with c[8]:
                if not locked and st.button("🗑️", key=f"del_{row_id}"):
                    to_delete_ids.add(row_id)

            updated_items.append({
                "_row_id": row_id,
                "물품명": name,
                "규격": spec if str(spec).strip() else "-",
                "단위": unit,
                "수량": qty,
                "금액": amount,
                "검사항목": insp_item,
                "검사결과": insp_result,
                "검사일자": insp_date.strftime("%Y-%m-%d"),
            })

        if to_delete_ids:
            st.session_state.receipt_items = [it for it in updated_items if it["_row_id"] not in to_delete_ids]
            st.rerun()
        else:
            st.session_state.receipt_items = updated_items

        if not locked and st.button("➕ 항목 추가"):
            st.session_state.receipt_items.append({
                "_row_id": _new_receipt_row_id(),
                "물품명": "",
                "규격": "-",
                "단위": "개",
                "수량": 1,
                "금액": 0,
                "검사항목": DEFAULT_INSPECTION_ITEM,
                "검사결과": DEFAULT_INSPECTION_RESULT,
                "검사일자": datetime.today().strftime("%Y-%m-%d"),
            })
            st.rerun()

        st.divider()
        st.subheader("📄 다운로드 미리보기")
        st.write(f"**상호명:** {store_name or '-'}")
        st.write(f"**구매날짜:** {purchase_date.strftime('%Y-%m-%d')}")
        preview_items = [
            {k: v for k, v in item.items() if k != "_row_id"}
            for item in st.session_state.receipt_items
        ]
        st.dataframe(preview_items, width="stretch", hide_index=True)

        st.divider()
        if not locked:
            confirmed = st.checkbox("✅ 이상없음을 확인했습니다. (체크 후 다운로드 가능)")
            if confirmed:
                st.session_state.confirmed = True
                st.rerun()
        else:
            if st.button("🔓 잠금 해제 (내용 수정하기)"):
                st.session_state.confirmed = False
                st.rerun()

        st.divider()
        if st.button("📄 DOCX/XLSX 파일 생성", type="primary", width="stretch", disabled=not st.session_state.confirmed):
            if not st.session_state.confirmed:
                st.warning("⚠️ '이상없음을 확인' 체크박스를 먼저 체크해주세요.")
            elif not report_title.strip():
                st.error("⚠️ 보고서 제목을 입력해주세요!")
            else:
                with st.spinner("문서를 생성하는 중..."):
                    try:
                        items = st.session_state.receipt_items
                        doc_buffer = build_receipt_docx(report_title, report_date, purchase_date, store_name, items)
                        xlsx_buffer = build_receipt_xlsx(report_title, report_date, purchase_date, store_name, items)

                        date_filename = report_date.strftime("%Y%m%d")
                        base_filename = f"{_safe_filename(report_title)}_{date_filename}_물품보고서"

                        st.success("✅ 파일이 생성되었어요!")
                        col_docx, col_xlsx = st.columns(2)
                        with col_docx:
                            st.download_button(
                                label="⬇️ DOCX 다운로드",
                                data=doc_buffer,
                                file_name=f"{base_filename}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                width="stretch",
                            )
                        with col_xlsx:
                            st.download_button(
                                label="⬇️ Excel 다운로드",
                                data=xlsx_buffer,
                                file_name=f"{base_filename}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                width="stretch",
                            )
                    except Exception as e:
                        st.error(f"오류: {str(e)}")
                        st.exception(e)
    elif uploaded_file:
        st.info("💡 '영수증 읽기'를 누르면 결과가 표시됩니다.")
